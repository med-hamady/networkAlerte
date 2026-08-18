"""Carte cartographique des sites — rendu en image, puis export Word.

Ce module dessine la **cartographie du parc** : les sites posés à leur vraie
place sur un fond de carte, et les backhauls tracés entre eux. C'est la vue
« plan mural » du réseau, celle qu'on imprime et qu'on fait circuler — à ne pas
confondre avec la page `/topology`, qui sert le même maillage sous forme de
graphe en couches et de carte Google interactive.

⚠️ **Le fond de carte est EMBARQUÉ, pas téléchargé.** Trois tuiles composées
(`data/maps/*.jpg` + `bounds.json`) sont livrées avec le code. Rien n'est
demandé à un serveur de tuiles au moment de l'export : le serveur de prod est
derrière un FortiGate et n'a pas d'accès Internet sortant garanti, et un export
qui dépend d'un CDN échouerait le jour où on en a besoin. Le prix à payer est
que le **cadrage est figé** — un site posé hors de la fenêtre d'une ville ne
peut pas être dessiné, et il est alors **nommé** dans le document plutôt
qu'escamoté (même règle que la vue carte : une carte qui omet un site sans le
dire se lit comme un réseau qui n'a pas ce site).

⚠️ **Deux populations, deux couleurs, deux sources.**

* **VERT — l'existant.** Les 17 sites de Nouakchott, lus dans la topologie à
  chaque export : coordonnées relevées (`site_locations`), liaisons mesurées.
  Un backhaul posé hier apparaît sans rien toucher ici.
* **ROUGE — les extensions PROGRAMMÉES**, pas encore installées : deux à
  Nouakchott, trois à Nouadhibou, deux à Rosso. Elles ne sont dans aucune de
  nos tables — c'est leur état NORMAL, elles n'existent pas encore — et vivent
  dans la constante `_PLANNED` ci-dessous. Leurs positions sont des
  **intentions**, pas des relevés, et ne sont **jamais écrites en base** : les y
  inscrire les ferait passer pour du parc installé, donc les ferait compter dans
  la capacité, pinguer, et alerter comme injoignables.

Le jour où un site programmé est posé et enrôlé dans UISP, il remonte tout seul
par le chemin normal, en vert — il ne reste plus qu'à retirer sa ligne de
`_PLANNED`.
"""

from __future__ import annotations

import io
import json
import logging
import math
from collections.abc import Iterable, Sequence
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont

logger = logging.getLogger(__name__)

_ASSETS = Path(__file__).resolve().parents[2] / "data" / "maps"

# ----------------------------------------------------------------- palette
# Les couleurs de trait viennent de la NATURE de la liaison, pas de son état :
# ce document est un plan de câblage, pas un tableau de supervision. L'état
# vivant se lit sur /topology, où il change toutes les minutes.
_INK = (15, 47, 94)
_GREEN = (21, 127, 60)      # backhaul radio de dorsale
_BLUE = (26, 95, 208)       # fibre / cuivre
_AMBER = (224, 138, 30)     # backhaul radio hors arbre = boucle de secours
_RED = (192, 39, 30)
_WHITE = (255, 255, 255)
_DEEP = (11, 36, 71)        # fond des cartouches « arrivée Internet »
_LABEL_BORDER = (195, 204, 217)

# Supersampling du calque de dessin. PIL ne lisse pas les formes : tracer les
# pastilles à 1× donne des cercles crénelés sur un fond de carte net. On dessine
# à 2× puis on réduit en LANCZOS — le coût est une image transitoire de ~80 Mo
# sur la plus grande planche, le temps d'une requête.
_SS = 2

_FONT_CANDIDATES = {
    True: (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        "C:/Windows/Fonts/arialbd.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
    ),
    False: (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
    ),
}


class MapAssetsError(RuntimeError):
    """Les fonds de carte livrés avec le code sont absents ou illisibles."""


@dataclass(frozen=True)
class Plate:
    """Une planche : une ville, sa fenêtre géographique, son fond de carte."""

    key: str
    title: str
    scale: float


_PLATES: tuple[Plate, ...] = (
    Plate("nouakchott", "Nouakchott", 1.0),
    Plate("nouadhibou", "Nouadhibou", 1.55),
    Plate("rosso", "Rosso", 1.2),
)

# La couleur d'une pastille dit son ÉTAT, jamais sa ville : vert = installé,
# rouge = programmé. Un code couleur par ville obligerait à lire une légende
# pour comprendre le message principal de la carte.
_PIN_INSTALLED = _GREEN
_PIN_PLANNED = _RED

# Extensions PROGRAMMÉES — pas encore installées (cf. en-tête du module).
# Positions approchées : elles situent l'intention dans la ville, ce ne sont pas
# des relevés. Rendues en ROUGE partout.
#
# ⚠️ Nouakchott en a aussi : la planche « vivante » porte donc les deux
# populations. C'est tout l'intérêt de la carte — voir d'un coup ce qui est
# debout et ce qui manque, sur le même plan.
#
# ⚠️ Aucune liaison n'est tracée depuis un site programmé de Nouakchott : son
# raccordement n'est pas arrêté. En dessiner une inventerait de la topologie,
# ce que ce module ne fait nulle part. Les liaisons de Nouadhibou et Rosso, en
# revanche, font partie du plan lui-même — ce sont des réseaux à créer d'un
# bloc, pas des ajouts à un maillage existant.
_PLANNED: dict[str, dict] = {
    "nouakchott": {
        "sites": [
            ("NKC-NORD", 18.1456, -15.9593),   # à côté d'AT2
            ("NKC-SUD", 18.0270, -15.9210),    # la zone laissée nue derrière VEL1
        ],
        "edges": [],
    },
    "nouadhibou": {
        "sites": [
            ("NDB-NORD", 20.9782, -17.0285),
            ("NDB-CENTRE", 20.9500, -17.0375),
            ("NDB-SUD", 20.9155, -17.0395),
        ],
        "edges": [
            ("NDB-NORD", "NDB-CENTRE"),
            ("NDB-CENTRE", "NDB-SUD"),
        ],
    },
    "rosso": {
        # Alignés à la MÊME latitude : les deux mâts se répondent d'ouest en
        # est le long du fleuve, pas du nord au sud.
        "sites": [
            ("RSO-NORD", 16.5155, -15.8100),
            ("RSO-SUD", 16.5155, -15.7950),
        ],
        "edges": [("RSO-NORD", "RSO-SUD")],
    },
}

# D'où vient Internet, et par quel site il entre dans la ville.
#
# C'est la seule chose que la carte affirme SANS la tenir d'une mesure : aucune
# de nos tables ne dit qu'un site est la tête de réseau (le lien amont n'est pas
# un data-link, cf. `TOPOLOGY_ROOT_SITE` — le contrôleur ignore lui aussi quel
# site fait face à l'amont). C'est donc un fait d'exploitation, écrit ici.
#
# `lat`/`lon` situent le CARTOUCHE, pas une installation : il est posé dans une
# zone vide du cadre, et recadré automatiquement s'il déborde.
_FEEDS: dict[str, list[dict]] = {
    "nouakchott": [
        {
            "target": "A2 HQ",
            "lat": 18.1150,
            "lon": -16.0330,
            "title": "INTERNET",
            "lines": ["arrivée nationale"],
        }
    ],
    "nouadhibou": [
        {
            "target": "NDB-CENTRE",
            "lat": 20.9930,
            "lon": -17.0640,
            "title": "INTERNET",
            "lines": ["depuis Nouakchott", "par câble fibre optique"],
        }
    ],
    "rosso": [
        {
            "target": "RSO-NORD",
            "lat": 16.5295,
            "lon": -15.8135,
            "title": "INTERNET",
            "lines": ["depuis Nouakchott", "par câble fibre optique"],
        }
    ],
}


# ------------------------------------------------------------- géométrie
def _mercator_y(lat: float) -> float:
    s = math.sin(math.radians(lat))
    return math.log((1 + s) / (1 - s)) / 2


def _load_bounds() -> dict[str, dict]:
    path = _ASSETS / "bounds.json"
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, ValueError) as exc:  # pragma: no cover - défaut d'install
        raise MapAssetsError(f"bounds.json illisible ({path}) : {exc}") from exc
    return {entry["name"]: entry for entry in raw}


def _project(bounds: dict, lat: float, lon: float) -> tuple[float, float]:
    """(lat, lon) → pixel dans le fond de carte, en projection Web Mercator."""
    y_north = _mercator_y(bounds["north"])
    y_south = _mercator_y(bounds["south"])
    x = (lon - bounds["west"]) / (bounds["east"] - bounds["west"]) * bounds["w"]
    y = (y_north - _mercator_y(lat)) / (y_north - y_south) * bounds["h"]
    return x, y


def _inside(bounds: dict, lat: float, lon: float) -> bool:
    return (bounds["south"] <= lat <= bounds["north"]
            and bounds["west"] <= lon <= bounds["east"])


# ------------------------------------------------------------- typographie
def _font(size: int, bold: bool = True) -> ImageFont.FreeTypeFont:
    for path in _FONT_CANDIDATES[bold]:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    # Aucune police vectorielle : plutôt qu'un rendu illisible en bitmap, on le
    # dit — l'image de la carte est le produit, une carte sans noms ne sert à
    # rien. L'image Docker installe `fonts-dejavu-core` pour cette raison.
    raise MapAssetsError(
        "Aucune police TrueType trouvée (DejaVu/Liberation/Arial). "
        "Installer fonts-dejavu-core dans l'image backend."
    )


# ------------------------------------------------------------- primitives
def _dashed_line(draw: ImageDraw.ImageDraw, p1, p2, color, width: float,
                 dash: float, gap: float) -> None:
    x1, y1 = p1
    x2, y2 = p2
    total = math.hypot(x2 - x1, y2 - y1)
    if total <= 0:
        return
    ux, uy = (x2 - x1) / total, (y2 - y1) / total
    pos = 0.0
    while pos < total:
        end = min(pos + dash, total)
        draw.line(
            [(x1 + ux * pos, y1 + uy * pos), (x1 + ux * end, y1 + uy * end)],
            fill=color, width=int(round(width)),
        )
        pos = end + gap


def _link(draw: ImageDraw.ImageDraw, p1, p2, color, width: float, dashed: bool,
          trim: float) -> None:
    """Trait entre deux pastilles, raccourci de `trim` à chaque bout.

    Doublé d'un liseré blanc dessous : sur un fond de carte clair chargé de
    routes jaunes, un trait de couleur seul se perd. Le liseré est plein même
    quand le trait est tireté — c'est lui qui porte la continuité du lien.
    """
    x1, y1 = p1
    x2, y2 = p2
    dist = math.hypot(x2 - x1, y2 - y1) or 1.0
    t = trim / dist
    a = (x1 + (x2 - x1) * t, y1 + (y2 - y1) * t)
    b = (x2 - (x2 - x1) * t, y2 - (y2 - y1) * t)
    draw.line([a, b], fill=(255, 255, 255, 190), width=int(round(width + 3.4)))
    if dashed:
        _dashed_line(draw, a, b, color, width, width * 3.1, width * 2.3)
    else:
        draw.line([a, b], fill=color, width=int(round(width)))


def _pin(draw: ImageDraw.ImageDraw, cx: float, cy: float, r: float,
         color: tuple[int, int, int]) -> None:
    """Pastille ronde + pointe + pictogramme de pylône, façon marqueur de carte."""
    ring = max(2.0, r * 0.16)
    draw.polygon(
        [(cx - r * 0.34, cy + r * 0.80), (cx, cy + r * 1.62), (cx + r * 0.34, cy + r * 0.80)],
        fill=color, outline=_WHITE, width=int(round(ring * 0.8)),
    )
    draw.ellipse([cx - r, cy - r, cx + r, cy + r], fill=color,
                 outline=_WHITE, width=int(round(ring)))

    g = r / 20.0          # le pictogramme est dessiné dans une boîte de 40 unités
    w = max(1.0, 2.4 * g)
    lw = int(round(w))
    draw.line([(cx - 5.4 * g, cy + 12 * g), (cx - 1.6 * g, cy - 8 * g)], fill=_WHITE, width=lw)
    draw.line([(cx + 5.4 * g, cy + 12 * g), (cx + 1.6 * g, cy - 8 * g)], fill=_WHITE, width=lw)
    draw.line([(cx - 3.8 * g, cy + 2 * g), (cx + 3.8 * g, cy + 2 * g)], fill=_WHITE, width=lw)
    draw.line([(cx - 4.9 * g, cy + 7.2 * g), (cx + 4.9 * g, cy + 7.2 * g)], fill=_WHITE, width=lw)
    for radius, span in ((8.0, 55), (13.0, 50)):
        box = [cx - radius * g, cy - (radius + 11) * g,
               cx + radius * g, cy + (radius - 11) * g]
        draw.arc(box, 180 + span, 360 - span, fill=_WHITE, width=lw)
        draw.arc(box, span, 180 - span, fill=_WHITE, width=lw)
    dot = 1.1 * g
    draw.ellipse([cx - dot, cy - 11 * g - dot, cx + dot, cy - 11 * g + dot], fill=_WHITE)


# Positions candidates d'un nom autour de sa pastille, dans l'ordre d'essai.
# Les quatre diagonales ne sont pas du luxe : un site très maillé (le siège en
# porte cinq liaisons plus l'arrivée Internet) n'a AUCUN des quatre côtés droits
# de libre, et le solveur choisissait alors le moins mauvais — en pratique, par-
# dessus le nom du voisin.
_LABEL_SIDES = ("e", "w", "s", "n", "ne", "nw", "se", "sw")


def _label_box(cx: float, cy: float, r: float, side: str, tw: float,
               th: float) -> tuple[float, float, float, float]:
    pad_x, pad_y = th * 0.55, th * 0.38
    w, h = tw + pad_x * 2, th + pad_y * 2
    gap = r * 0.55
    left, right = cx - r - gap - w, cx + r + gap
    above, below = cy - r - gap - h, cy + r * 1.7
    positions = {
        "e": (right, cy - h / 2),
        "w": (left, cy - h / 2),
        "s": (cx - w / 2, below),
        "n": (cx - w / 2, above),
        "ne": (cx + r * 0.45, above),
        "nw": (cx - r * 0.45 - w, above),
        "se": (cx + r * 0.45, below),
        "sw": (cx - r * 0.45 - w, below),
    }
    x, y = positions[side]
    return x, y, x + w, y + h


def _overlap(a, b) -> float:
    dx = min(a[2], b[2]) - max(a[0], b[0])
    dy = min(a[3], b[3]) - max(a[1], b[1])
    return dx * dy if dx > 0 and dy > 0 else 0.0


def _segment_samples(segments: Iterable[tuple[tuple[float, float], tuple[float, float]]],
                     step: float) -> list[tuple[float, float]]:
    """Échantillonne les traits de liaison en points, pour le test de recouvrement.

    Un test segment/rectangle exact serait plus élégant ; l'échantillonnage est
    suffisant ici (on cherche à savoir si un nom se pose SUR un trait, pas à
    mesurer de combien) et ne peut pas se tromper de signe.
    """
    points: list[tuple[float, float]] = []
    for (x1, y1), (x2, y2) in segments:
        length = math.hypot(x2 - x1, y2 - y1)
        count = max(2, int(length / max(step, 1.0)))
        for i in range(count + 1):
            t = i / count
            points.append((x1 + (x2 - x1) * t, y1 + (y2 - y1) * t))
    return points


def _place_labels(items: Sequence[tuple[str, float, float]], r: float,
                  font: ImageFont.FreeTypeFont, draw: ImageDraw.ImageDraw,
                  canvas: tuple[float, float],
                  segments: Sequence[tuple[tuple[float, float], tuple[float, float]]],
                  obstacles: Sequence[tuple] = (),
                  ) -> list[tuple[str, tuple]]:
    """Choisit de quel côté de sa pastille chaque nom se pose.

    Placement glouton, essayé dans l'ordre est / ouest / sud / nord, noté par
    ce qu'il recouvrirait : les autres noms déjà posés, les pastilles, **les
    traits de liaison**, et le débord hors cadre. Volontairement AUTOMATIQUE et
    non une table de réglages par site : un backhaul posé sur un nouveau site
    doit apparaître à l'export suivant sans qu'on ait à toucher au code.

    ⚠️ Les traits comptent, et c'est le point qui change tout : sans eux, tous
    les noms se posent à l'est et masquent précisément les liaisons que la carte
    est là pour montrer. Ils sont pénalisés plus lourdement qu'un simple
    chevauchement de surface — un nom qui mord sur un autre nom se lit encore,
    un nom posé sur un backhaul l'efface.

    L'ordre de parcours est figé (nord → sud) pour que deux exports des mêmes
    données rendent exactement la même image.
    """
    width, height = canvas
    pins = [(x - r, y - r, x + r, y + r * 1.7) for _, x, y in items]
    samples = _segment_samples(segments, r * 0.35)
    # Les cartouches d'arrivée Internet sont posés AVANT les noms et ne bougent
    # plus : ils entrent donc dans le calcul comme des noms déjà placés.
    placed: list[tuple] = list(obstacles)
    out: list[tuple[str, tuple]] = []
    for name, x, y in items:
        left, top, right, bottom = draw.textbbox((0, 0), name, font=font)
        tw, th = right - left, bottom - top
        best, best_cost = None, None
        for side in _LABEL_SIDES:
            box = _label_box(x, y, r, side, tw, th)
            area = max((box[2] - box[0]) * (box[3] - box[1]), 1.0)
            hits = sum(
                1 for px, py in samples
                if box[0] <= px <= box[2] and box[1] <= py <= box[3]
            )
            # Trois gênes RAMENÉES À LA MÊME ÉCHELLE (une fraction de la surface
            # du nom), sans quoi les poids ne sont pas comparables et le terme
            # brut le plus grand décide seul.
            #
            # ⚠️ Le recouvrement d'un AUTRE NOM pèse le plus lourd : deux noms
            # superposés sont illisibles, alors qu'un nom posé sur un trait
            # masque un trait qu'on devine encore. C'était l'inverse avant, et
            # « A2 HQ » se posait sur « A2 NR1 » pour éviter ses liaisons.
            cost = (
                4.0 * sum(_overlap(box, other) for other in placed) / area
                + 1.5 * sum(_overlap(box, pin) for pin in pins) / area
                + 2.0 * min(1.0, hits / 8.0)
            )
            if box[0] < 0 or box[1] < 0 or box[2] > width or box[3] > height:
                cost += 1e6
            if best_cost is None or cost < best_cost:
                best, best_cost = box, cost
        placed.append(best)
        out.append((name, best))
    return out


def _globe(draw: ImageDraw.ImageDraw, cx: float, cy: float, r: float,
           color: tuple[int, int, int]) -> None:
    """Petit globe : ce qui fait lire le cartouche comme « Internet » d'un coup."""
    lw = max(1, int(round(r * 0.16)))
    draw.ellipse([cx - r, cy - r, cx + r, cy + r], outline=color, width=lw)
    draw.line([(cx - r, cy), (cx + r, cy)], fill=color, width=lw)
    draw.arc([cx - r * 0.48, cy - r, cx + r * 0.48, cy + r], 0, 360, fill=color, width=lw)


def _clip_to_box(box: tuple[float, float, float, float],
                 target: tuple[float, float]) -> tuple[float, float]:
    """Point du bord du cartouche vers la cible — le trait ne le traverse pas."""
    cx, cy = (box[0] + box[2]) / 2, (box[1] + box[3]) / 2
    tx, ty = target
    dx, dy = tx - cx, ty - cy
    if dx == 0 and dy == 0:
        return cx, cy
    half_w, half_h = (box[2] - box[0]) / 2, (box[3] - box[1]) / 2
    scale = min(
        half_w / abs(dx) if dx else float("inf"),
        half_h / abs(dy) if dy else float("inf"),
    )
    return cx + dx * scale, cy + dy * scale


def _feed_geometry(feeds: Iterable[dict], bounds: dict, scale: float,
                   draw: ImageDraw.ImageDraw, anchors: dict[str, tuple[float, float]],
                   ) -> list[dict]:
    """Place les cartouches d'arrivée Internet et calcule leur trait.

    Le cartouche est **recadré dans la planche** s'il déborde : sa position
    n'est qu'une zone vide visée à la main, alors que sa taille dépend du texte
    et de la police trouvée sur la machine. Sans ce recadrage, un mot de plus
    dans le libellé sortirait la moitié du cartouche hors de l'image.
    """
    title_font = _font(int(round(40 * scale)))
    line_font = _font(int(round(27 * scale)), bold=False)
    pad = 20 * scale
    gap = 9 * scale
    placed: list[dict] = []

    for feed in feeds:
        target = anchors.get(feed["target"])
        if target is None:
            continue  # le site d'entrée n'est pas dessiné : rien à raccorder

        globe_r = 19 * scale
        rows = [(feed["title"], title_font)] + [(t, line_font) for t in feed["lines"]]
        sizes = []
        for text, font in rows:
            left, top, right, bottom = draw.textbbox((0, 0), text, font=font)
            sizes.append((right - left, bottom - top))
        width = max(w for w, _h in sizes) + globe_r * 2 + gap
        height = sum(h for _w, h in sizes) + gap * (len(rows) - 1)

        x, y = _project(bounds, feed["lat"], feed["lon"])
        cx, cy = x * _SS, y * _SS
        box = [cx - width / 2 - pad, cy - height / 2 - pad,
               cx + width / 2 + pad, cy + height / 2 + pad]

        margin = 12 * scale
        max_x, max_y = bounds["w"] * _SS, bounds["h"] * _SS
        shift_x = max(margin - box[0], 0) - max(box[2] - (max_x - margin), 0)
        shift_y = max(margin - box[1], 0) - max(box[3] - (max_y - margin), 0)
        box = [box[0] + shift_x, box[1] + shift_y, box[2] + shift_x, box[3] + shift_y]

        placed.append({
            "box": tuple(box),
            "rows": rows,
            "sizes": sizes,
            "globe_r": globe_r,
            "gap": gap,
            "pad": pad,
            "anchor": _clip_to_box(tuple(box), target),
            "target": target,
        })
    return placed


def _draw_feed_box(draw: ImageDraw.ImageDraw, feed: dict) -> None:
    box = feed["box"]
    pad, gap, globe_r = feed["pad"], feed["gap"], feed["globe_r"]
    draw.rounded_rectangle(box, radius=pad * 0.8, fill=_DEEP,
                           outline=_WHITE, width=max(2, int(round(pad * 0.16))))

    text_left = box[0] + pad + globe_r * 2 + gap
    total = sum(h for _w, h in feed["sizes"]) + gap * (len(feed["sizes"]) - 1)
    y = (box[1] + box[3]) / 2 - total / 2
    _globe(draw, box[0] + pad + globe_r, (box[1] + box[3]) / 2, globe_r, _WHITE)
    for (text, font), (_w, h) in zip(feed["rows"], feed["sizes"], strict=True):
        draw.text((text_left, y + h / 2), text, font=font, fill=_WHITE, anchor="lm")
        y += h + gap


def _compass(draw: ImageDraw.ImageDraw, bounds: dict, scale: float) -> None:
    r = 34 * scale
    cx = bounds["w"] * _SS - (r + 26 * scale)
    cy = bounds["h"] * _SS - (r + 26 * scale)
    draw.ellipse([cx - r, cy - r, cx + r, cy + r], fill=(255, 255, 255, 235),
                 outline=_LABEL_BORDER, width=int(round(2 * scale)))
    draw.polygon(
        [(cx, cy - r * 0.34), (cx + r * 0.26, cy + r * 0.58),
         (cx, cy + r * 0.33), (cx - r * 0.26, cy + r * 0.58)],
        fill=_INK,
    )
    # Le « N » tient DANS le disque : posé plus haut, il dépasse du cercle et la
    # rose se lit comme un glyphe cassé.
    font = _font(int(round(15 * scale)))
    draw.text((cx, cy - r * 0.42), "N", font=font, fill=_INK, anchor="ms")


def _attribution(draw: ImageDraw.ImageDraw, bounds: dict, scale: float) -> None:
    """Mention de source de l'imagerie — exigée par la licence du fournisseur.

    Elle est dessinée DANS l'image : le document Word n'a pas d'autre endroit
    où elle survivrait à un copier-coller de la carte. Le texte vient de
    `bounds.json`, donc du script qui a téléchargé les tuiles — un changement
    de fournisseur d'imagerie ne peut pas oublier de changer la mention.
    """
    text = bounds.get("attribution") or "© Esri"
    font = _font(int(round(19 * scale)), bold=False)
    x, y = 14 * scale, bounds["h"] * _SS - 14 * scale
    left, top, right, bottom = draw.textbbox((x, y), text, font=font, anchor="ls")
    pad = 6 * scale
    draw.rounded_rectangle(
        [left - pad, top - pad, right + pad, bottom + pad],
        radius=5 * scale, fill=(255, 255, 255, 225),
    )
    draw.text((x, y), text, font=font, fill=(91, 106, 125), anchor="ls")


# ------------------------------------------------------------- rendu
def _plate_data(plate: Plate, topo: dict, bounds: dict
                ) -> tuple[list[tuple[str, float, float, bool]],
                           list[tuple[str, str, bool, bool]], list[str]]:
    """Sites et liaisons d'une planche, plus les sites non traçables.

    Un site est rendu `(nom, lat, lon, programmé)`. L'installé de Nouakchott
    sort de la topologie — donc de la base — et les extensions programmées de
    `_PLANNED` ; les deux se retrouvent sur la même planche.
    """
    sites: list[tuple[str, float, float, bool]] = []
    missing: list[str] = []

    for site in topo.get("sites", []) if plate.key == "nouakchott" else []:
        lat, lon = site.get("latitude"), site.get("longitude")
        name = site.get("site", "")
        if lat is None or lon is None:
            missing.append(f"{name} (position inconnue)")
            continue
        if not _inside(bounds, lat, lon):
            missing.append(f"{name} (hors du cadrage de la carte)")
            continue
        sites.append((" ".join(name.split()), lat, lon, False))

    installed = {name for name, _lat, _lon, _p in sites}
    planned = _PLANNED.get(plate.key, {"sites": [], "edges": []})
    for name, lat, lon in planned["sites"]:
        # Un site programmé qui existe désormais en base est déjà dessiné en
        # vert : on ne le redouble pas en rouge. C'est ce qui rend le retrait de
        # sa ligne de `_PLANNED` facultatif le jour de sa mise en service.
        if name in installed:
            continue
        sites.append((name, lat, lon, True))

    drawable = {name for name, _lat, _lon, _p in sites}
    edges: list[tuple[str, str, bool, bool]] = []
    for edge in topo.get("edges", []) if plate.key == "nouakchott" else []:
        a = " ".join(str(edge.get("site_a", "")).split())
        b = " ".join(str(edge.get("site_b", "")).split())
        if a not in drawable or b not in drawable:
            continue
        wired = edge.get("medium") == "wired"
        edges.append((a, b, not wired, bool(edge.get("is_tree_edge", True))))

    for a, b in planned["edges"]:
        if a in drawable and b in drawable:
            edges.append((a, b, True, True))

    return sites, edges, missing


def render_plate(plate: Plate, topo: dict, bounds_by_key: dict[str, dict]
                 ) -> tuple[bytes, list[str]]:
    """Dessine une planche et rend le JPEG plus la liste des sites non traçables."""
    bounds = bounds_by_key.get(plate.key)
    if bounds is None:
        raise MapAssetsError(f"Fenêtre géographique absente pour « {plate.key} »")
    try:
        base = Image.open(_ASSETS / f"{plate.key}.jpg").convert("RGBA")
    except OSError as exc:  # pragma: no cover - défaut d'installation
        raise MapAssetsError(f"Fond de carte « {plate.key} » illisible : {exc}") from exc

    sites, edges, missing = _plate_data(plate, topo, bounds)

    scale = plate.scale
    r = 32 * scale * _SS
    width = 7.2 * scale * _SS
    font = _font(int(round(35 * scale * _SS)))

    overlay = Image.new("RGBA", (bounds["w"] * _SS, bounds["h"] * _SS), (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay)

    # La pastille est posée AU-DESSUS du point, sa pointe sur les coordonnées :
    # c'est la pointe qui désigne le lieu, pas le centre du disque.
    anchors: dict[str, tuple[float, float]] = {}
    for name, lat, lon, _planned in sites:
        x, y = _project(bounds, lat, lon)
        anchors[name] = (x * _SS, y * _SS - 1.25 * r)

    segments: list[tuple[tuple[float, float], tuple[float, float]]] = []
    for a, b, dashed, tree in edges:
        color = _AMBER if (dashed and not tree) else (_GREEN if dashed else _BLUE)
        _link(draw, anchors[a], anchors[b], color, width, dashed, r + width)
        segments.append((anchors[a], anchors[b]))

    # L'arrivée Internet se raccorde comme une fibre : trait plein, même bleu.
    # Le trait passe SOUS les pastilles et le cartouche AU-DESSUS des noms —
    # c'est l'élément qui répond à « d'où vient Internet ici ».
    feeds = _feed_geometry(_FEEDS.get(plate.key, ()), bounds, scale * _SS, draw, anchors)
    for feed in feeds:
        _link(draw, feed["anchor"], feed["target"], _BLUE, width, False, r + width)
        segments.append((feed["anchor"], feed["target"]))

    ordered = sorted(sites, key=lambda s: -s[1])
    for name, _lat, _lon, planned in ordered:
        cx, cy = anchors[name]
        _pin(draw, cx, cy, r, _PIN_PLANNED if planned else _PIN_INSTALLED)

    items = [(name, *anchors[name]) for name, _lat, _lon, _planned in ordered]
    canvas = (bounds["w"] * _SS, bounds["h"] * _SS)
    obstacles = [feed["box"] for feed in feeds]
    for name, box in _place_labels(items, r, font, draw, canvas, segments, obstacles):
        draw.rounded_rectangle(box, radius=(box[3] - box[1]) * 0.30, fill=_WHITE,
                               outline=_LABEL_BORDER, width=int(round(2 * scale * _SS / 2)))
        draw.text(((box[0] + box[2]) / 2, (box[1] + box[3]) / 2), name,
                  font=font, fill=_INK, anchor="mm")

    for feed in feeds:
        _draw_feed_box(draw, feed)

    _compass(draw, bounds, scale * _SS)
    _attribution(draw, bounds, scale * _SS)

    overlay = overlay.resize((bounds["w"], bounds["h"]), Image.LANCZOS)
    composed = Image.alpha_composite(base, overlay).convert("RGB")

    # JPEG et pas PNG : la planche finie est une PHOTO (imagerie satellite) que
    # le sans-perte fait peser dix fois plus — 9,9 Mo contre 1,6 Mo mesurés sur
    # Nouakchott, pour un document Word qui dépassait alors 13 Mo. La qualité 88
    # garde les étiquettes nettes ; ce sont elles, pas les toits, qu'on lit.
    buffer = io.BytesIO()
    composed.save(buffer, format="JPEG", quality=88, optimize=True, progressive=True)
    return buffer.getvalue(), missing


def render_plates(topo: dict) -> tuple[list[tuple[Plate, bytes, int, int]], list[str]]:
    """Rend les trois planches.

    Retourne `(planche, png, installés, programmés)` par ville, plus les sites
    qu'aucune planche n'a pu dessiner. Les deux compteurs restent SÉPARÉS
    jusqu'au document : les additionner annoncerait comme parc en service des
    mâts qui ne sont pas montés.
    """
    bounds_by_key = _load_bounds()
    out: list[tuple[Plate, bytes, int, int]] = []
    missing: list[str] = []
    for plate in _PLATES:
        png, plate_missing = render_plate(plate, topo, bounds_by_key)
        sites = _plate_data(plate, topo, bounds_by_key[plate.key])[0]
        installed = sum(1 for _n, _la, _lo, planned in sites if not planned)
        out.append((plate, png, installed, len(sites) - installed))
        missing.extend(plate_missing)
    return out, missing


# ------------------------------------------------------------- export Word
def _sync_label(topo: dict) -> str:
    raw = str(topo.get("synced_at") or "")[:10]
    parts = raw.split("-")
    return f"{parts[2]}/{parts[1]}/{parts[0]}" if len(parts) == 3 else "-"


def _network_paragraphs(plates) -> tuple[str, str]:
    """Les deux paragraphes d'ouverture : l'existant, puis ce qui est programmé.

    Les chiffres sont CALCULÉS à partir des planches, jamais écrits en dur — un
    document dont le texte contredit sa propre carte ne serait plus lu.
    """
    installed = sum(count for _plate, _png, count, _planned in plates)
    planned = sum(count for _plate, _png, _installed, count in plates)
    cities = [
        f"{plate.title} ({count})"
        for plate, _png, _installed, count in plates
        if count
    ]

    existing = (
        f"Notre réseau actuel couvre Nouakchott avec une bonne capacité : "
        f"{installed} sites d'infrastructure sont en service, raccordés au siège "
        f"par des dorsales fibre et un maillage de faisceaux radio, avec des "
        f"boucles de secours qui permettent à la plupart des sites d'atteindre "
        f"Internet par plusieurs chemins. L'ensemble est supervisé en continu : "
        f"disponibilité, qualité radio et charge de chaque liaison."
    )
    extension = (
        f"Les sites en rouge sont les extensions programmées, non encore "
        f"installées : {planned} nouveaux sites — {', '.join(cities)} — "
        f"destinés à compléter la couverture là où elle manque et à augmenter "
        f"la capacité du réseau. Leur emplacement est une intention de "
        f"déploiement ; ils n'apparaissent pas encore dans la supervision."
    )
    return existing, extension


def build_topology_docx(topo: dict) -> bytes:
    """Assemble le document Word : une ville par page, en A4 portrait.

    Une planche par page et pas les trois sur une : réduites au tiers d'une
    page, les étiquettes de site deviennent illisibles — or c'est pour les lire
    qu'on imprime ce document.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    plates, missing = render_plates(topo)

    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    for attr in ("left_margin", "right_margin"):
        setattr(section, attr, Cm(1.5))
    section.top_margin, section.bottom_margin = Cm(1.4), Cm(1.2)
    usable = Cm(18.0)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Cartographie des sites A2 Holding")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(*_INK)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Câblage du {_sync_label(topo)}")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(91, 106, 125)

    # Les deux paragraphes ouvrent le document, AVANT la première carte : ils
    # disent quoi lire dans les couleurs. Les découvrir après les planches
    # obligerait à revenir en arrière pour comprendre le rouge.
    existing, extension = _network_paragraphs(plates)
    for text, color in ((existing, _INK), (extension, _RED)):
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para_run = para.add_run(text)
        para_run.font.size = Pt(10.5)
        para_run.font.color.rgb = RGBColor(*color)

    for index, (plate, png, installed, planned) in enumerate(plates):
        if index:
            document.add_page_break()
        heading = document.add_paragraph()
        parts = []
        if installed:
            parts.append(f"{installed} site{'s' if installed > 1 else ''} en service")
        if planned:
            parts.append(f"{planned} programmé{'s' if planned > 1 else ''}")
        head_run = heading.add_run(f"{plate.title} — {' · '.join(parts)}")
        head_run.bold = True
        head_run.font.size = Pt(14)
        head_run.font.color.rgb = RGBColor(*_INK)
        picture = document.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.add_run().add_picture(io.BytesIO(png), width=usable)

    if missing:
        # Un site qu'on n'a pas pu dessiner est NOMMÉ. Une carte silencieuse sur
        # ce qu'elle omet se lit comme un réseau qui n'a pas ces sites.
        note = document.add_paragraph()
        note_run = note.add_run(
            "Sites non représentés : " + " · ".join(missing)
        )
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(192, 39, 30)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
